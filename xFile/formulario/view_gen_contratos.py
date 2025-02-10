from django.conf import settings
from django.contrib.auth.mixins import LoginRequiredMixin
from django.urls import reverse_lazy
from django.contrib import messages
from .models import (
                    Datos,
                    Contratos,)

from django.http import HttpResponse
from django.shortcuts import render, redirect
from rest_framework import status

from datetime import datetime
from docx import Document
import os

def genContratosTotales(request):
    registros = Datos.objects.select_related('representante').all()
    if registros.exists():
        try:
            for registro in registros:
                print(f"Serial: {registro.representante.cedula.upper()}, Empresa: {registro.representante.representante.upper()}")
            #generar_docx_desde_plantilla(registros)
            messages.success(request, "Los contratos se generaron exitosamente.")
        except Exception as e:
            messages.error(request, f"Ocurrió un error al generar los contratos: {str(e)}")
    else:
        messages.warning(request, "No se encontraron registros para generar contratos.")

    # Redirigir a otra vista
    return redirect('datos_list')  # Reemplaza 'nombre_de_la_vista' con el nombre de la vista a la que deseas redirigir

def genContratosUnico(request, pk):
    registro = Datos.objects.get(pk=pk) 
    
    if registro.exists():
        try:
            generar_docx_desde_plantilla(registro)
            messages.success(request, "El contrato se generó exitosamente.")
        except Exception as e:
            messages.error(request, f"Ocurrió un error al generar el contrato: {str(e)}")
    else:
        messages.warning(request, "No se encontraron registros para generar el contrato.")

    # Redirigir a otra vista
    return redirect('datos_list')  


def generar_docx_desde_plantilla(registros):
    """
    Genera un archivo .docx basado en una plantilla, reemplazando los placeholders.

    Args:
        reemplazos (dict): Diccionario con los placeholders y sus valores.
    """
    ruta_plantilla = os.path.join(settings.MEDIA_ROOT, 'plantillas', 'modelo_venta_a_plazo.docx')
    ruta_salida = os.path.join(settings.MEDIA_ROOT, 'contratos')
    
    # Crear la carpeta de salida si no existe
    os.makedirs(ruta_salida, exist_ok=True) 

    for contrato, registro in enumerate(registros, start=1):
        # Crear una copia de la plantilla
        doc = Document(ruta_plantilla)
    
        # Reemplazar placeholders en cada párrafo
        for parrafo in doc.paragraphs:
            # Reemplazar texto
            if "{serial_cliente}" in parrafo.text:
                reemplazar_texto(parrafo, "{serial_cliente}", registro.serial_cliente.upper())  # Nombre en mayúsculas
            if "{empresa_r}" in parrafo.text:
                reemplazar_texto(parrafo, "{empresa_r}", registro.representante.representante.upper())  # Fecha actual
    
        # Reemplazar placeholders en tablas (opcional)
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for placeholder, valor in registro.__dict__.items():
                        if placeholder in celda.text:
                            celda.text = celda.text.replace(placeholder, str(valor))
                            
        fecha = datetime.now().strftime('%Y%m%d')
        nombre_archivo = f"{registro.serial_cliente}_{fecha}.docx"  # Usa el serial_cliente para el nombre
        ruta_archivo = os.path.join(ruta_salida, nombre_archivo)
        doc.save(ruta_archivo)
            
        # Guardar la ruta en el modelo Contratos
        contrato_obj = Contratos.objects.create(
                serial_cliente=registro,  # Asignar el objeto Datos
                url_contrato=nombre_archivo  # Guardar la ruta del archivo
        )


def reemplazar_texto(parrafo, buscar, reemplazar):
    if buscar in parrafo.text:
        inline = parrafo.runs
        for i in inline:
            if buscar in i.text:
                i.text = i.text.replace(buscar, reemplazar)